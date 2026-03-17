from docx import Document

# Doc 1: Master Guide
doc1 = Document()
doc1.add_heading('SIMILAN DIGITAL AGENCY', 0)
doc1.add_heading('Bot System Master Guide', 1)
doc1.add_paragraph('Version 1.0 | March 15, 2026')

bots = [
    ("Content Manager", "Social copy, captions, hashtags", "No health claims, respect platform cadence"),
    ("Visual Designer", "Graphics, layouts, templates", "3 variations per brief, brand colors"),
    ("Paid Ads Bot", "Meta/Google campaigns, optimization", "ROAS >1.2:1, budget guardrails"),
    ("Community Manager", "Comments, DMs, engagement", "<2h response time"),
    ("Analytics Bot", "KPIs, dashboards, reporting", "Real-time data, weekly digest"),
    ("Copywriter Bot", "Long-form: emails, blogs", "Persuasive, authentic, strategic"),
    ("Video Editor Bot", "Short-form Reels, TikToks", "15-30s, hook in 0-3s"),
    ("Influencer/Outreach Bot", "Partnerships, collaborations", "Phase 3+ only, Dan approval"),
    ("Project Manager Bot", "Workflow orchestration, Asana", "Daily status, escalate blockers"),
    ("Security & Compliance Bot", "24/7 threat monitoring", "12-point security checklist"),
]

for name, role, constraints in bots:
    doc1.add_heading(name, 2)
    doc1.add_paragraph(f"Role: {role}", style='List Bullet')
    doc1.add_paragraph(f"Constraints: {constraints}", style='List Bullet')

doc1.add_heading('Launch Action Checklist', 1)
for item in ["Approve all 10 bot roles", "Assign Asana project", "Set up Telegram bot", "Configure Discord", "Create API keys", "Test bot communication", "Verify approval gates", "Set up monitoring", "Document escalation", "Run Week 1 dry-run"]:
    doc1.add_paragraph(item, style='List Bullet')

doc1.save('/Users/pandamac/.openclaw/workspace/SIMILAN_BOT_SYSTEM_MASTER_GUIDE.docx')

# Doc 2: Implementation
doc2 = Document()
doc2.add_heading('DAN IMPLEMENTATION CHECKLIST', 0)
doc2.add_heading('Your Setup Work - March 15-31, 2026', 1)
doc2.add_heading('External Tools Matrix', 2)

for tool, tier, purpose in [("Asana", "Free/Pro", "PM"), ("Meta Business Suite", "Free", "Social"), ("GA4", "Free", "Analytics"), ("Mailchimp/Klaviyo", "Starter+", "Email"), ("Telegram", "Free", "Alerts")]:
    doc2.add_paragraph(f"{tool} ({tier}): {purpose}", style='List Bullet')

doc2.add_heading('Dan Action Items (TODAY)', 2)
for item in ["Get Meta credentials", "Confirm GA4 IDs", "Pick email platform", "Create Asana workspace", "Generate Telegram token", "Secure API keys", "Approve bot roles"]:
    doc2.add_paragraph(f"☐ {item}", style='List Bullet')

doc2.save('/Users/pandamac/.openclaw/workspace/DAN_IMPLEMENTATION_CHECKLIST.docx')

print("Done - both docs created")
